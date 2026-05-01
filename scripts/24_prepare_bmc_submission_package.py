from __future__ import annotations

import subprocess
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission" / "bmc_medical_genomics_2026-05-01"
MANUSCRIPT_MD = ROOT / "reports" / "manuscript" / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_TARGET_DRAFT.md"
TITLE_PAGE_MD = SUBMISSION_DIR / "TITLE_PAGE_DRAFT.md"
COVER_LETTER_MD = SUBMISSION_DIR / "COVER_LETTER_DRAFT.md"
SUBMISSION_BODY_MD = SUBMISSION_DIR / "_generated_submission_manuscript_body.md"

MAIN_DOCX = SUBMISSION_DIR / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_SUBMISSION_DRAFT.docx"
COVER_DOCX = SUBMISSION_DIR / "COVER_LETTER_DRAFT.docx"


def run_pandoc(inputs: list[Path], output: Path) -> None:
    cmd = [
        "pandoc",
        *[str(path) for path in inputs],
        "--standalone",
        "-o",
        str(output),
    ]
    subprocess.run(cmd, check=True)


def write_submission_body() -> Path:
    text = MANUSCRIPT_MD.read_text(encoding="utf-8")
    text = re.sub(r"^# .+?\n+", "", text, count=1)
    text = re.sub(r"^Target journal draft:.*?\n+", "", text, count=1, flags=re.M)
    text = re.sub(r"^## Running Title\n.*?(?=^## Abstract\n)", "", text, count=1, flags=re.S | re.M)
    SUBMISSION_BODY_MD.write_text(text.lstrip(), encoding="utf-8")
    return SUBMISSION_BODY_MD


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run("1")
    paragraph.add_run()._r.append(end)


def enable_line_numbers(section) -> None:
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:lnNumType")):
        sect_pr.remove(old)

    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def format_main_manuscript(path: Path) -> None:
    doc = Document(path)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        enable_line_numbers(section)
        footer = section.footer
        footer.paragraphs[0].clear()
        add_page_number(footer.paragraphs[0])

    for style_name in ["Normal", "Body Text"]:
        if style_name in [style.name for style in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(12)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 2.0
        paragraph.paragraph_format.space_after = Pt(6)

    doc.save(path)


def format_cover_letter(path: Path) -> None:
    doc = Document(path)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(8)

    doc.save(path)


def main() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    submission_body = write_submission_body()

    run_pandoc([TITLE_PAGE_MD, submission_body], MAIN_DOCX)
    run_pandoc([COVER_LETTER_MD], COVER_DOCX)

    format_main_manuscript(MAIN_DOCX)
    format_cover_letter(COVER_DOCX)

    print(MAIN_DOCX)
    print(COVER_DOCX)


if __name__ == "__main__":
    main()
