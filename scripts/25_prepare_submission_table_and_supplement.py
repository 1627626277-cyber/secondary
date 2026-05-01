from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission" / "bmc_medical_genomics_2026-05-01"

TABLE_MD = SUBMISSION_DIR / "TABLE1_CROSS_COHORT_EVIDENCE_DRAFT.md"
TABLE_DOCX = SUBMISSION_DIR / "TABLE1_CROSS_COHORT_EVIDENCE_DRAFT.docx"

SUPP_MD = SUBMISSION_DIR / "SUPPLEMENTARY_METHODS_REPRODUCIBILITY_DRAFT.md"
SUPP_DOCX = SUBMISSION_DIR / "SUPPLEMENTARY_METHODS_REPRODUCIBILITY_DRAFT.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))


def set_cell_margins(cell, margin_dxa: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ["top", "start", "bottom", "end"]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_dxa))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def parse_markdown_table(path: Path) -> tuple[str, list[str], list[list[str]], str]:
    text = path.read_text(encoding="utf-8")
    title = next(line[2:].strip() for line in text.splitlines() if line.startswith("# "))
    note_match = re.search(r"^Note:\s*(.+)$", text, flags=re.M)
    note = note_match.group(1).strip() if note_match else ""
    table_lines = [line for line in text.splitlines() if line.startswith("|")]
    header = [cell.strip() for cell in table_lines[0].strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        rows.append([cell.strip() for cell in line.strip("|").split("|")])
    return title, header, rows, note


def format_paragraph_font(paragraph, size: float = 9, bold: bool = False, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def build_table_docx() -> None:
    title, header, rows, note = parse_markdown_table(TABLE_MD)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True

    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    widths = [1.35, 0.7, 1.95, 3.5, 2.35]
    header_cells = table.rows[0].cells
    repeat_table_header(table.rows[0])
    for idx, cell in enumerate(header_cells):
        cell.text = header[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "1F4E79")
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell, 110)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph_font(para, size=8.5, bold=True, color=RGBColor(255, 255, 255))

    for data_row in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell.text = data_row[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, 100)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.05
                format_paragraph_font(para, size=8.2)

    note_para = doc.add_paragraph()
    note_para.paragraph_format.space_before = Pt(8)
    note_run = note_para.add_run("Note: " + note)
    note_run.font.name = "Arial"
    note_run.font.size = Pt(8.5)

    doc.save(TABLE_DOCX)


def run_pandoc(input_path: Path, output_path: Path) -> None:
    subprocess.run(["pandoc", str(input_path), "--standalone", "-o", str(output_path)], check=True)


def format_supplement_docx() -> None:
    run_pandoc(SUPP_MD, SUPP_DOCX)
    doc = Document(SUPP_DOCX)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    for style_name in ["Normal", "Body Text"]:
        if style_name in [style.name for style in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(10)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.name = "Arial"

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, 80)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.05
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8)

    doc.save(SUPP_DOCX)


def main() -> None:
    build_table_docx()
    format_supplement_docx()
    print(TABLE_DOCX)
    print(SUPP_DOCX)


if __name__ == "__main__":
    main()
