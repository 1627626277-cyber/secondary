from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "submission" / "bmc_medical_genomics_2026-05-01"
MANUSCRIPT_MD = ROOT / "reports" / "manuscript" / "MANUSCRIPT_BMC_MEDICAL_GENOMICS_TARGET_DRAFT.md"
TITLE_PAGE_MD = SUBMISSION_DIR / "TITLE_PAGE_DRAFT.md"
TABLE1_MD = SUBMISSION_DIR / "TABLE1_CROSS_COHORT_EVIDENCE_DRAFT.md"
LEGENDS_MD = ROOT / "reports" / "manuscript" / "FIGURE_LEGENDS_DRAFT.md"
FIGURE_DIR = ROOT / "analysis" / "manuscript_figures"

OUT_DOCX = SUBMISSION_DIR / "MANUSCRIPT_CELLS_STYLE_READING_VERSION.docx"
OUT_MD = SUBMISSION_DIR / "MANUSCRIPT_CELLS_STYLE_READING_VERSION_SOURCE.md"


FIGURES = {
    "Figure 1": FIGURE_DIR / "fig1_study_design_evidence_chain.png",
    "Figure 2": FIGURE_DIR / "fig2_spatial_plasma_secretory_discovery.png",
    "Figure 3": FIGURE_DIR / "fig3_scrna_plasma_secretory_localization.png",
    "Figure 4": FIGURE_DIR / "fig4_geo_bulk_clinical_support.png",
    "Figure 5": FIGURE_DIR / "fig5_commppass_os_iss_validation.png",
    "Figure 6": FIGURE_DIR / "fig6_gse299193_xenium_spatial_validation.png",
}


HEADING_COLOR = RGBColor(0x14, 0x5A, 0x72)
ACCENT = RGBColor(0x24, 0x7C, 0x93)
BODY = RGBColor(0x20, 0x20, 0x20)


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def first_h1(text: str) -> str:
    match = re.search(r"^#\s+(.+)$", text, re.M)
    if not match:
        raise RuntimeError("No H1 title found")
    return match.group(1).strip()


def section(text: str, name: str) -> str:
    pattern = rf"^##\s+{re.escape(name)}\s*$"
    start = re.search(pattern, text, re.M)
    if not start:
        return ""
    next_h2 = re.search(r"^##\s+", text[start.end() :], re.M)
    end = start.end() + next_h2.start() if next_h2 else len(text)
    return text[start.end() : end].strip()


def parse_h3_sections(text: str) -> list[tuple[str, str]]:
    matches = list(re.finditer(r"^###\s+(.+)$", text, re.M))
    if not matches:
        return [("", text.strip())]
    sections: list[tuple[str, str]] = []
    leading = text[: matches[0].start()].strip()
    if leading:
        sections.append(("", leading))
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        sections.append((match.group(1).strip(), text[start:end].strip()))
    return sections


def parse_figure_legends(text: str) -> dict[str, str]:
    matches = list(re.finditer(r"^##\s+(Figure\s+\d+)\s*$", text, re.M))
    legends: dict[str, str] = {}
    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        legends[match.group(1)] = text[start:end].strip()
    return legends


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, margin: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run_markup(paragraph, text: str, size: float | None = None, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = bold_default
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        run = paragraph.add_run(part.replace("`", ""))
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        run.font.color.rgb = BODY


def add_paragraph(doc: Document, text: str, style: str = "Body Text", space_after: float = 6) -> None:
    text = clean_inline(text)
    if not text:
        return
    para = doc.add_paragraph(style=style)
    add_run_markup(para, text, 10.2)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.05
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_markdown_blocks(doc: Document, text: str, level_base: int = 2) -> None:
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        if buffer:
            add_paragraph(doc, " ".join(buffer))
            buffer = []

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("### "):
            flush()
            add_heading(doc, clean_inline(line[4:]), level_base + 1)
        elif line.startswith("## "):
            flush()
            add_heading(doc, clean_inline(line[3:]), level_base)
        elif line.startswith("- "):
            flush()
            para = doc.add_paragraph(style="Body Text")
            para.style = "List Bullet"
            add_run_markup(para, clean_inline(line[2:]), 10.0)
            para.paragraph_format.space_after = Pt(3)
        else:
            buffer.append(line)
    flush()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.style = f"Heading {min(level, 3)}"
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = HEADING_COLOR
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11.5)
    para.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    para.paragraph_format.space_after = Pt(5)


def add_title_page(doc: Document, title: str, title_md: str) -> None:
    meta = {
        "Authors": "",
        "Affiliations": "",
        "ORCID": "",
        "Keywords": "",
    }
    current = None
    for line in title_md.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            current = stripped[3:]
            continue
        if current in meta and stripped:
            meta[current] = (meta[current] + " " + clean_inline(stripped)).strip()

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = label.add_run("Article")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    label.paragraph_format.space_after = Pt(10)

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x11, 0x37, 0x4A)
    para.paragraph_format.space_after = Pt(8)

    authors = doc.add_paragraph()
    authors.add_run("Zhuang Jiang").bold = True
    authors.runs[0].font.name = "Times New Roman"
    authors.runs[0].font.size = Pt(11)
    authors.paragraph_format.space_after = Pt(4)

    aff = doc.add_paragraph()
    aff.add_run("1. Guangdong University of Petrochemical Technology (GDUPT), 139 Guandu 2nd Road, Maoming 525000, Guangdong, China").italic = True
    aff.runs[0].font.name = "Times New Roman"
    aff.runs[0].font.size = Pt(9)
    aff.paragraph_format.space_after = Pt(2)

    orcid = doc.add_paragraph()
    orcid.add_run("ORCID: https://orcid.org/0009-0007-4388-5901")
    orcid.runs[0].font.name = "Times New Roman"
    orcid.runs[0].font.size = Pt(9)
    orcid.paragraph_format.space_after = Pt(10)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "6AA8B6")
    borders.append(bottom)
    p_pr.append(borders)


def add_abstract(doc: Document, abstract: str, keywords: str) -> None:
    add_heading(doc, "Abstract", 1)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F8FA")
    set_cell_margins(cell, 140)
    for para in cell.paragraphs:
        para.text = ""

    for heading, body in parse_h3_sections(abstract):
        if not heading:
            continue
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f"{heading}: ")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.6)
        run.font.color.rgb = HEADING_COLOR
        add_run_markup(para, " ".join(body.split()), 9.6)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    key_para = doc.add_paragraph()
    key_para.paragraph_format.space_before = Pt(8)
    key_para.paragraph_format.space_after = Pt(12)
    r = key_para.add_run("Keywords: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9.8)
    add_run_markup(key_para, " ".join(keywords.split()), 9.8)


def add_table1(doc: Document, table_md: str) -> None:
    add_heading(doc, "Table 1. Cross-cohort evidence chain and manuscript claim boundary", 3)
    rows = []
    for line in table_md.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|") or set(stripped.replace("|", "").strip()) <= {"-", ":"}:
            continue
        cells = [clean_inline(x) for x in stripped.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.15, 0.55, 1.20, 2.35, 1.35]
    for r_idx, row in enumerate(rows):
        if r_idx == 0:
            set_repeat_table_header(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 80)
            if c_idx < len(widths):
                cell.width = Inches(widths[c_idx])
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (1,) else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(7.2 if r_idx else 7.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "145A72")
            else:
                run.font.color.rgb = BODY
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "F5F9FA")

    note = re.search(r"Note:\s*(.+)$", table_md, re.S)
    if note:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run("Note: ")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        add_run_markup(para, " ".join(note.group(1).split()), 8)


def add_figure(doc: Document, fig_name: str, legend: str) -> None:
    path = FIGURES[fig_name]
    if not path.exists():
        raise FileNotFoundError(path)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(6.65))
    para.paragraph_format.space_after = Pt(4)

    first, *rest = legend.splitlines()
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.space_after = Pt(10)
    lead = caption.add_run(f"{fig_name}. ")
    lead.bold = True
    lead.font.name = "Arial"
    lead.font.size = Pt(8.8)
    add_run_markup(caption, first.strip(), 8.8, bold_default=True)
    if rest:
        body = " ".join(line.strip() for line in rest if line.strip())
        caption.add_run(" ")
        add_run_markup(caption, body, 8.5)


def add_references(doc: Document, refs: str) -> None:
    add_heading(doc, "References", 1)
    for para_text in refs.splitlines():
        text = para_text.strip()
        if not text:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        add_run_markup(para, text, 8.2)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.70)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.70)
    section.right_margin = Inches(0.70)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = BODY

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body.font.size = Pt(10.2)


def add_page_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        footer.add_run()._r.append(begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        footer.add_run()._r.append(instr)

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        footer.add_run()._r.append(separate)

        number = footer.add_run("1")
        number.font.size = Pt(8)
        number.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        footer.add_run()._r.append(end)


def add_main_text(doc: Document, manuscript: str, legends: dict[str, str]) -> None:
    background = section(manuscript, "Background")
    methods = section(manuscript, "Methods")
    results = section(manuscript, "Results")
    discussion = section(manuscript, "Discussion")
    conclusions = section(manuscript, "Conclusions")
    declarations = section(manuscript, "Declarations")
    references = section(manuscript, "References")

    add_heading(doc, "1. Introduction", 1)
    add_markdown_blocks(doc, background)

    add_heading(doc, "2. Materials and Methods", 1)
    method_idx = 1
    for title, body in parse_h3_sections(methods):
        if title:
            add_heading(doc, f"2.{method_idx}. {title}", 2)
            method_idx += 1
        add_markdown_blocks(doc, body)

    add_heading(doc, "3. Results", 1)
    result_sections = {title: body for title, body in parse_h3_sections(results) if title}
    result_order = [
        "Cross-cohort design defined a bounded evidence chain",
        "Spatial transcriptomics identified the plasma-secretory program (Fig. 2)",
        "Single-cell data localized the axis to plasma-cell compartments (Fig. 3)",
        "External bulk cohorts supported a clinical subtype module (Fig. 4)",
        "CoMMpass/GDC linked the axis to OS, ISS, and molecular risk (Fig. 5)",
        "Adjusted models supported covariate-adjusted OS association (Fig. 5D)",
        "Independent Xenium data reproduced the program-level spatial signal (Fig. 6)",
    ]
    result_idx = 1
    for title in result_order:
        body = result_sections.get(title)
        if not body:
            continue
        add_heading(doc, f"3.{result_idx}. {re.sub(r'\\s+\\(Fig\\. .*?\\)$', '', title)}", 2)
        add_markdown_blocks(doc, body)
        if title.startswith("Cross-cohort"):
            add_figure(doc, "Figure 1", legends["Figure 1"])
            add_table1(doc, read(TABLE1_MD))
        elif title.startswith("Spatial transcriptomics"):
            add_figure(doc, "Figure 2", legends["Figure 2"])
        elif title.startswith("Single-cell"):
            add_figure(doc, "Figure 3", legends["Figure 3"])
        elif title.startswith("External bulk"):
            add_figure(doc, "Figure 4", legends["Figure 4"])
        elif title.startswith("Adjusted models"):
            add_figure(doc, "Figure 5", legends["Figure 5"])
        elif title.startswith("Independent Xenium"):
            add_figure(doc, "Figure 6", legends["Figure 6"])
        result_idx += 1

    add_heading(doc, "4. Discussion", 1)
    add_markdown_blocks(doc, discussion)

    add_heading(doc, "5. Conclusions", 1)
    add_markdown_blocks(doc, conclusions)

    add_heading(doc, "Declarations", 1)
    add_markdown_blocks(doc, declarations)

    add_references(doc, references)


def build_source_markdown(manuscript: str, legends: dict[str, str]) -> None:
    text = manuscript
    for name, image in FIGURES.items():
        text += f"\n\n![{name}]({image.as_posix()})\n\n{name}. {legends.get(name, '').splitlines()[0]}\n"
    OUT_MD.write_text(text, encoding="utf-8")


def main() -> None:
    manuscript = read(MANUSCRIPT_MD)
    title_md = read(TITLE_PAGE_MD)
    legends = parse_figure_legends(read(LEGENDS_MD))

    doc = Document()
    setup_styles(doc)
    add_title_page(doc, first_h1(manuscript), title_md)
    add_abstract(doc, section(manuscript, "Abstract"), section(manuscript, "Keywords"))
    add_main_text(doc, manuscript, legends)
    add_page_footer(doc)
    build_source_markdown(manuscript, legends)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
